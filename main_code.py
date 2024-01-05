import requests
from bs4 import BeautifulSoup
import os
import time
from docx import Document
from docx.shared import Pt
from collections import Counter

vuln_count = 0
vuln_list = []


def extract_forms(webpage):
    soup = BeautifulSoup(webpage, "html.parser")
    return soup.find_all("form")


def url_to_html(url):
    page = requests.get(url)
    return page.text


def parse_forms(forms):
    result = []
    for form in forms:
        try:
            form_data = {}
            form_data['action'] = form['action'] if form.has_attr('action') else None
            form_data['method'] = form['method'] if form.has_attr('method') else None
            form_data['inputs'] = [inputs for inputs in form.find_all('input')]
            result.append(form_data)
        except: 
            continue
    return result


def fuzz(payloads, input_elem, url, method, delay):
    global vuln_count, vuln_list
    for script in payloads:
        print(f"[{vuln_count}] Testing script: {script}")

        if method.lower() == "get":
            response = requests.get(url, params={input_elem['name']: script}).text
        else:
            response = requests.post(url, data={input_elem['name']: script}).text

        if script in response:
            vuln_count += 1
            vuln_list.append([input_elem, script])
            print(f"\033[91m[+] Vulnerability found!\033[0m\n")

        if delay is not None:
            time.sleep(delay)


def read_payloads_from_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return [line.strip() for line in file]


def save_result_to_file(result_file, vuln_list):
    count = 1
    base_name, extension = os.path.splitext(result_file)
    new_result_file = result_file

    while os.path.exists(new_result_file):
        new_result_file = f"{base_name} ({count}){extension}"
        count += 1

    input_strings = [vuln[0] for vuln in vuln_list]
    frequency = Counter(input_strings)
    most_common_string = frequency.most_common(1)[0][0]

    with open(new_result_file, 'w', encoding='utf-8') as file:
        file.write("================= RESULT =================\n")
        file.write(f"{len(vuln_list)} Vulnerabilities Found!\n\n")

        for count, vuln in enumerate(vuln_list):
            file.write(f"[{count+1}] {vuln[0]}\n\t{vuln[1]}\n\n")

    return most_common_string


def PrintResult():
    global vuln_list, vuln_count
    print("================= RESULT =================\n")
    print(f"\033[91m{vuln_count} Vulnerabilities Found!\033[0m\n")

    input_strings = [vuln[0] for vuln in vuln_list]
    frequency = Counter(input_strings)
    most_common_string = frequency.most_common(1)[0][0]

    for count, vuln in enumerate(vuln_list):
        print(f"[{count+1}] {vuln[0]}\n\t\033[91m{vuln[1]}\033[0m\n")

    return most_common_string


def create_docx_with_hello(file_path, url, most_common_string):
    count = 1
    base_name, extension = os.path.splitext(file_path)
    new_file_path = file_path

    while os.path.exists(new_file_path):
        new_file_path = f"{base_name} ({count}){extension}"
        count += 1

    doc = Document()

    paragraph1 = doc.add_paragraph(f"\"{url}\"에서의 취약점 스캔 결과로 Reflected XSS 취약점이 발견되었습니다.")
    
    run1 = paragraph1.runs[0]
    run1.font.bold = True

    doc.add_paragraph("    Reflected XSS 공격은 악성 스크립트가 포함된 URL을 사용자가 클릭하도록 유도하여 URL을 클릭하면 클라이언트를 공격하는 방식으로 동작합니다. 일반적으로 서버에 검색 내용을 입력하면, 검색 결과가 있는 경우에는 결과 값을 사용자에게 전달하지만, 서버에서 정확한 결과가 없는 경우 서버는 브라우저에 입력한 값을 그대로 HTML 문서에 포함하여 응답합니다. 이 경우 HTML 페이지에 포함된 악성 스크립트가 브라우저에서 실행이 됩니다.")    
    doc.add_paragraph(f"    현재 스캐닝을 진행한 웹사이트에서는 {most_common_string} 부분에서 가장 많은 취약점이 발견되었습니다. 이는 세선ID 유출, 시스템 관리자 권한 탈취, 악성코드 유포 등 보다 확대된 피해를 낳을 수도 있습니다. 그러므로 아래와 같은 조치를 시행하는 것을 권장합니다.\n")

    doc.add_paragraph("1. 입・출력 값 검증 및 무효화")
    doc.add_paragraph("    XSS 취약점을 근본적으로 제거하기 위해서는 스크립트 등 해킹에 사용될 수 있는 코딩에 사용되는 입력 및 출력 값에 대해서 검증하고 무효화시켜야 합니다. 입력 값에 대한 유효성 검사는 데이터가 입력되기 전에 가능하면, 입력 데이터에 대한 길이, 문자, 형식 및 사업적 규칙 유효성을 검사해야 합니다.\n  하지만 애플리케이션 개발자가 많은 태그의 입력 문자를 검증하기 위해 코딩 시 일일이 작업하는 것은 많은 노력과 자원이 소모됩니다. 또한 인코딩 방식을 통해 방어기술을 무력화할 수 있으므로 애플리케이션 개발자 직접 모두 처리하는 것은 근본적으로 불가능합니다. 그러므로 입・출력 값을 자동적으로 검증해주는 라이브러리를 사용하면 좀 더 효과적으로 대응할 수 있습니다.")

    doc.add_paragraph("2. 보안 라이브러리")
    doc.add_paragraph("1) AntiXSS")
    doc.add_paragraph("    AntiXSS 라이브러리는 마이크로소프트사에서 개발한 공개용 XSS 취약점 예방 라이브러리입니다. AntiXSS는 입력 값을 검증하여 서버로 악성 스크립트로 입력되지 못하는 기능과 위험한 문자를 인코딩하는 함수를 제공합니다.")

    doc.add_paragraph("2) OWASP ESAPI 라이브러리")
    doc.add_paragraph("    OWASP는 포괄적인 애플리케이션 보안을 위해 웹 응용 취약점을 대응할 수 있는 오픈소스 ESAPI 라이브러리를 개발하여 제공하고 있습니다. ESAPI에는 총 14개의 API가 있으며, 이 중 XSS 취약점을 예방하기 위해 API는 validator와 encoder가 있습니다. validator는 입력 값을 필터링하는 기능을 하고 있으며, encoder는 출력 값을 인코딩 및 디코딩 기능을 가지고 있습니다. 이 라이브러리는 자바, PHP, .NET, ASP, 자바스크립트 및 파이썬 등 다양한 애플리케이션 개발 언어를 지원하고 있어 보다 편리한 사용이 가능합니다.")
    
    doc.add_paragraph("\n본 보고서는 한국인터넷진흥원의 KISA Report 중, 크로스사이트스크립팅(XSS) 공격종류 및 대응방법의 내용으로 기술하였음을 알려드립니다.")

    paragraph4 = doc.paragraphs[3]
    for run in paragraph4.runs:
        run.font.size = Pt(14)

    paragraph6 = doc.paragraphs[5]
    for run in paragraph6.runs:
        run.font.size = Pt(14)

    doc.save(new_file_path)


input_str = input("> ")
inputSplit = input_str.split()
url = inputSplit[0]
options = inputSplit[1:]
delay = None 

script_path = os.path.abspath(__file__)
dirname = os.path.dirname(script_path)
payloads_file = os.path.join(dirname, "xss_vectors.txt")


if "-t" in options:
    delay_idx = options.index("-t") + 1
    delay = int(options[delay_idx])

if "-c" in options:
    file_idx = options.index("-c") + 1
    payloads_file = options[file_idx]

dirname = os.path.dirname(os.path.abspath(__file__))
webpage = url_to_html(url)
forms = extract_forms(webpage)
parsed_forms = parse_forms(forms)

print("\n===== Start scanning vulnerabilities =====\n")

payloads = read_payloads_from_file(payloads_file)

for form in parsed_forms:
    try:
        method = form['method'].lower() if form['method'] else None
        inputs = form['inputs']
        for input_elem in inputs:
            fuzz(payloads, input_elem, url, method, delay)
    except:
        continue

print("\n" + "="*42 + "\n\n")

most_common_string = ""
while True:
    flag = input("Do you want to save the results as a text file? [Y/N]: ")
    if flag.lower() == 'y':
        result_file = os.path.join(dirname, "FuzzResult.txt")
        most_common_string = save_result_to_file(result_file, vuln_list)
        break
    elif flag.lower() == 'n':
        most_common_string = PrintResult()
        break
    else:
        print("Invalid input. Please enter 'Y' or 'N'")
    
while True:
    flag = input("Do you want to save the XSS_Vulnerability_Report? [Y/N]: ")
    if flag.lower() == 'y':
        result_file = os.path.join(dirname, "XSS_Vulnerability_Report.docx")
        create_docx_with_hello(result_file, url, most_common_string)
        break
    elif flag.lower() == 'n':
        break
    else:
        print("Invalid input. Please enter 'Y' or 'N'")
